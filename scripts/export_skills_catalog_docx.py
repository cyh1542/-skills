#!/usr/bin/env python3
"""Export all Olist skills: tools list + I/O definitions to a Word document."""
from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(1)

SKILLS_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = SKILLS_ROOT / "Olist_Skills_工具清单与输入输出定义.docx"


def _h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def _p(doc, text="", bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return para


def _bullets(doc, items: list[str]):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item).font.size = Pt(11)


def _table(doc, headers: list[str], rows: list[list[str]]):
    if not rows:
        _p(doc, "（无）")
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci in range(len(headers)):
            val = row[ci] if ci < len(row) else ""
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph("")


def build_catalog() -> Document:
    doc = Document()

    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Olist 电商分析 Skills")
    r.bold = True
    r.font.size = Pt(24)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("工具清单与输入输出定义").font.size = Pt(16)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"文档版本：1.0  |  生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph("")

    _h(doc, "文档说明", 1)
    _bullets(doc, [
        "本文档描述 Olist 巴西电商数据分析 Skills 的工具、输入、输出及失败条件。",
        "安装位置：项目 .cursor/skills/ 或 ~/.cursor/skills/（见仓库 README.md）。",
        "数据口径统一参见 olist-data-reference/reference.md。",
        "推荐执行顺序：weekly-ops-report → report-assets-summary → export_report_docx（Word 交付）。",
    ])

    # 总体关系
    _h(doc, "一、Skill 总览与依赖关系", 1)
    _table(doc, ["层级", "Skill 名称", "角色"], [
        ["主流程", "weekly-ops-report", "多表分析 → 周报 + charts + tables"],
        ["子模块", "delay-diagnosis", "履约延迟诊断"],
        ["子模块", "review-issue-locator", "差评归因定位"],
        ["子模块", "seller-risk-scan", "卖家风险扫描"],
        ["子模块", "category-experience-audit", "品类体验审计"],
        ["后处理", "report-assets-summary", "基于 charts/tables 生成 summary"],
        ["参考", "olist-data-reference", "表结构、Join、指标口径（不单独执行）"],
    ])

    _h(doc, "二、共享工具与脚本", 1)
    _table(doc, ["工具/脚本", "路径", "用途"], [
        ["文件读取器", "Cursor Read / pandas.read_csv", "加载 CSV"],
        ["Python DataFrame", "pandas", "清洗、Join、聚合、派生字段"],
        ["图表生成器", "matplotlib + seaborn", "PNG 趋势图/条形图/散点图"],
        ["中文字体", "weekly-ops-report/scripts/chart_zh_font.py", "setup_chinese_chart()、FontProperties"],
        ["表校验", "weekly-ops-report/scripts/validate_tables.py", "检查关键 CSV 与字段"],
        ["周报生成", "weekly-ops-report/SKILL.md（Agent 编排）", "主 skill 分析流程"],
        ["总结生成", "report-assets-summary/scripts/summarize_report_assets.py", "summary.md + summary.csv"],
        ["Word 导出", "report-assets-summary/scripts/export_report_docx.py", "巡检报告 .docx"],
        ["Word 模板", "report-assets-summary/templates/weekly_inspection_report_template.docx", "周报版式模板"],
        ["网页搜索", "WebSearch", "可选，节假日/外部解释"],
        ["报告输出器", "Write / 文件写入", "Markdown、CSV、DOCX"],
    ])

    skills = [
        skill_weekly_ops_report(),
        skill_delay_diagnosis(),
        skill_review_issue_locator(),
        skill_seller_risk_scan(),
        skill_category_experience_audit(),
        skill_report_assets_summary(),
        skill_olist_data_reference(),
    ]

    for i, block in enumerate(skills, start=1):
        doc.add_page_break()
        _h(doc, f"三.{i}  {block['title']}", 1)
        _p(doc, f"Skill 标识：{block['name']}", bold=True)
        _p(doc, block["desc"])
        doc.add_paragraph("")

        _h(doc, "适用场景", 2)
        _bullets(doc, block["scenarios"])

        _h(doc, "输入", 2)
        if block.get("input_table"):
            _table(doc, block["input_table"][0], block["input_table"][1])
        if block.get("input_params"):
            _p(doc, "参数：")
            _bullets(doc, block["input_params"])

        _h(doc, "工具", 2)
        _table(doc, ["工具", "用途"], block["tools"])

        _h(doc, "输出", 2)
        _table(doc, ["产出", "说明"], block["outputs"])

        if block.get("output_files"):
            _h(doc, "输出文件清单", 3)
            _table(doc, ["文件/目录", "说明"], block["output_files"])

        _h(doc, "失败条件（中止）", 2)
        _bullets(doc, block["failures"])

        _h(doc, "人工接管点", 2)
        _bullets(doc, block["manual"])

        if block.get("related"):
            _h(doc, "关联 Skill", 2)
            _bullets(doc, block["related"])

    # 端到端
    doc.add_page_break()
    _h(doc, "四、端到端输入输出一览", 1)
    _table(doc, ["阶段", "输入", "输出"], [
        ["① 原始分析", "数据目录 *.csv（Olist 各表）", "reports/weekly_YYYYMMDD/"],
        ["① 产出物", "—", "report.md、charts/、tables/、子模块 .md"],
        ["② 总结", "reports/weekly_*/（charts + tables）", "summary.csv（必选）、summary.md"],
        ["③ Word", "同上 + summary.csv", "*_report.docx、可选模板 .docx"],
    ])

    _h(doc, "五、推荐命令", 1)
    cmds = [
        "表校验：python weekly-ops-report/scripts/validate_tables.py <data_dir>",
        "总结：python report-assets-summary/scripts/summarize_report_assets.py <report_dir>",
        "Word 报告：python report-assets-summary/scripts/export_report_docx.py <report_dir>",
        "Word 模板：python report-assets-summary/scripts/create_word_template.py",
        "生成本目录文档：python scripts/export_skills_catalog_docx.py",
        "周报：在 Cursor Agent 中调用 weekly-ops-report skill，或按 SKILL.md 自行编排分析脚本",
    ]
    for c in cmds:
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(c)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    return doc


def skill_weekly_ops_report():
    return {
        "title": "weekly-ops-report（经营巡检周报）",
        "name": "weekly-ops-report",
        "desc": "基于 Olist 多表数据自动生成经营巡检周报：订单履约、延迟、差评、卖家风险、品类体验。",
        "scenarios": [
            "每周经营例会前的订单履约与客户体验分析",
            "需要同时交付文字结论 + 数据表 + 图表",
            "用户询问「近一周经营情况」「生成巡检周报」",
        ],
        "input_table": [
            ["逻辑名", "数据集", "必选"],
            ["orders", "olist_orders_dataset", "是"],
            ["order_items", "olist_order_items_dataset", "是"],
            ["reviews", "olist_order_reviews_dataset", "是"],
            ["payments", "olist_order_payments_dataset", "推荐"],
            ["products", "olist_products_dataset", "推荐"],
            ["customers", "olist_customers_dataset", "可选（区域）"],
            ["sellers", "olist_sellers_dataset", "可选"],
        ],
        "input_params": [
            "data_dir：CSV 所在目录",
            "start_date / end_date：分析窗口（默认订单量最高自然周或近 7 天）",
            "output_dir：默认 reports/weekly_{date}/",
        ],
        "tools": [
            ["validate_tables.py", "校验表与字段"],
            ["chart_zh_font.py", "图表中文（须 fontproperties）"],
            ["pandas", "多表 Join、指标计算"],
            ["matplotlib / seaborn", "≥3 张 PNG"],
            ["Write", "report.md"],
            ["WebSearch", "可选，外部因素解释"],
        ],
        "outputs": [
            ["report.md", "主周报"],
            ["charts/*.png", "≥3 张图表"],
            ["tables/*.csv", "≥3 张明细表"],
            ["delay_diagnosis.md 等", "子模块摘要（可选）"],
        ],
        "output_files": [
            ["weekly_delay_trend.csv", "周度指标趋势"],
            ["seller_risk_list.csv", "卖家风险清单"],
            ["problem_categories.csv", "问题品类"],
            ["delay_by_seller.csv", "卖家延迟 TOP"],
            ["delay_by_category.csv", "品类延迟 TOP"],
            ["delay_by_region.csv", "区域延迟 TOP"],
            ["review_by_delay_bins.csv", "延迟分箱 vs 差评率"],
        ],
        "failures": [
            "缺 orders / order_items / reviews",
            "时间字段无法解析",
            "主明细连接丢失率 > 20%",
            "分析窗口有效订单 < 30",
        ],
        "manual": [
            "处罚/下架/赔偿仅输出「待人工确认」",
            "环比异常 >50% 需业务确认",
        ],
        "related": [
            "子模块：delay-diagnosis、review-issue-locator、seller-risk-scan、category-experience-audit",
            "下游：report-assets-summary",
        ],
    }


def skill_delay_diagnosis():
    return {
        "title": "delay-diagnosis（延迟履约诊断）",
        "name": "delay-diagnosis",
        "desc": "诊断近周期履约是否恶化，延迟集中在哪些卖家/品类/区域。",
        "scenarios": ["履约是否恶化？", "延迟交付集中在哪些卖家、品类或区域？", "周报履约模块"],
        "input_table": [
            ["文件", "必选", "关键字段"],
            ["orders.csv", "是", "order_status, 购买/送达/预计送达时间"],
            ["order_items.csv", "是", "order_id, seller_id, product_id"],
            ["products.csv", "否", "product_category_name"],
            ["sellers/customers.csv", "否", "区域维度"],
        ],
        "input_params": ["window_days=7", "compare_previous=true"],
        "tools": [
            ["pandas", "延迟天数、延迟率计算"],
            ["matplotlib/seaborn + chart_zh_font", "趋势图、TOP 条形图"],
            ["文件写入", "CSV + Markdown"],
        ],
        "outputs": [
            ["delay_diagnosis.md", "结论与建议"],
            ["charts/delay_trend.png", "延迟率趋势"],
            ["charts/delay_top_sellers.png", "卖家 TOP"],
            ["tables/delay_by_*.csv", "卖家/品类/区域明细"],
        ],
        "failures": [
            "缺送达时间字段",
            "delivered 订单 < 50",
            "明细匹配丢失 > 20%",
        ],
        "manual": ["暂停合作/换物流须人工确认", "单卖家订单 <20 不进 TOP"],
        "related": ["被 weekly-ops-report 编排"],
    }


def skill_review_issue_locator():
    return {
        "title": "review-issue-locator（差评归因定位）",
        "name": "review-issue-locator",
        "desc": "分析差评与物流时效、价格、运费、品类、卖家的关联。",
        "scenarios": ["差评是否与物流/价格/运费/品类/卖家有关？", "体验复盘、客诉分析"],
        "input_table": [
            ["文件", "必选"],
            ["orders, order_items, reviews", "是"],
            ["products, payments", "否"],
        ],
        "input_params": ["bad_review_threshold=2（review_score ≤ 2 为差评）"],
        "tools": [
            ["pandas", "宽表、分箱、分组聚合"],
            ["matplotlib/seaborn + chart_zh_font", "差评率对比图"],
            ["scipy/pandas corr", "可选相关性"],
        ],
        "outputs": [
            ["review_attribution.md", "归因结论"],
            ["charts/bad_review_by_delay_bins.png", "延迟分箱 vs 差评率"],
            ["charts/bad_review_top_categories.png", "高差评 TOP"],
            ["tables/review_by_factor.csv", "因素对比表"],
        ],
        "failures": ["无 reviews", "有评论订单 < 100", "无法关联 orders > 15%"],
        "manual": ["不自动发道歉/补偿", "处罚认定需人工复核"],
        "related": ["被 weekly-ops-report 编排"],
    }


def skill_seller_risk_scan():
    return {
        "title": "seller-risk-scan（卖家风险扫描）",
        "name": "seller-risk-scan",
        "desc": "识别销量高但延迟率高、差评多、取消率高的卖家，输出风险得分。",
        "scenarios": ["哪些卖家在放大平台风险？", "周报风险卖家模块"],
        "input_table": [
            ["文件", "说明"],
            ["orders, order_items, reviews", "必选"],
            ["sellers", "推荐"],
        ],
        "input_params": ["min_orders=50", "top_n=20"],
        "tools": [
            ["pandas", "聚合 + 风险分（0-100）"],
            ["matplotlib + chart_zh_font", "风险散点图"],
        ],
        "outputs": [
            ["seller_risk_list.csv", "seller_id, 延迟率, 差评率, risk_score, risk_level"],
            ["seller_risk_summary.md", "摘要"],
            ["charts/seller_risk_scatter.png", "散点图"],
        ],
        "failures": ["无 seller_id", "满足 min_orders 卖家 < 5"],
        "manual": ["禁止自动下架/罚款/封店", "阈值调整需业务确认"],
        "related": ["被 weekly-ops-report 编排"],
    }


def skill_category_experience_audit():
    return {
        "title": "category-experience-audit（品类体验审计）",
        "name": "category-experience-audit",
        "desc": "识别规模大但延迟高、差评多的「问题类目」。",
        "scenarios": ["哪些品类规模大但体验差？", "品类结构优化"],
        "input_table": [
            ["文件", "必选"],
            ["orders, order_items, products, reviews", "是"],
        ],
        "input_params": ["min_category_orders=30"],
        "tools": [
            ["pandas", "品类聚合"],
            ["matplotlib + chart_zh_font", "健康矩阵图"],
        ],
        "outputs": [
            ["category_health.md", "结论"],
            ["tables/problem_categories.csv", "问题品类及 priority"],
            ["charts/category_health_matrix.png", "矩阵图"],
        ],
        "failures": ["商品映射失败 > 15%", "有效品类 < 5"],
        "manual": ["品类下架/收缩需负责人确认"],
        "related": ["被 weekly-ops-report 编排"],
    }


def skill_report_assets_summary():
    return {
        "title": "report-assets-summary（巡检结果总结）",
        "name": "report-assets-summary",
        "desc": "根据 charts/ 与 tables/ 归纳总结，输出 summary.csv（必选）与 summary.md。",
        "scenarios": [
            "weekly-ops-report 已产出 charts/tables，需二次总结",
            "用户要求生成 summary.csv",
        ],
        "input_table": [
            ["参数", "说明"],
            ["report_dir", "含 charts/ 和/或 tables/"],
            ["--md", "默认 summary.md"],
            ["--csv", "默认 summary.csv（必选）"],
        ],
        "tools": [
            ["summarize_report_assets.py", "生成 MD + CSV"],
            ["export_report_docx.py", "导出 Word 报告"],
            ["create_word_template.py", "生成 Word 版式模板"],
            ["Read（图片）", "可选，补充图表视觉结论"],
        ],
        "outputs": [
            ["summary.csv", "结构化总结（必选）"],
            ["summary.md", "可读版（推荐）"],
            ["*_report.docx", "Word 报告（可选）"],
        ],
        "output_files": [
            ["summary.csv", "9 列：类别/子类/指标名称/指标值/对象ID/说明/数据来源/优先级/需人工确认"],
        ],
        "failures": ["目录不存在", "charts 与 tables 皆空", "CSV 编码错误"],
        "manual": ["处罚类建议标注待确认", "图文冲突以 CSV 为准"],
        "related": ["上游：weekly-ops-report"],
    }


def skill_olist_data_reference():
    return {
        "title": "olist-data-reference（数据参考）",
        "name": "olist-data-reference（非执行类）",
        "desc": "统一表名、Join 路径、派生指标口径；供其它 skill 引用。",
        "scenarios": ["确认字段名、Join 关系、延迟率/差评率计算口径"],
        "input_table": [
            ["类型", "内容"],
            ["输入", "无（只读参考）"],
            ["输出", "无"],
        ],
        "tools": [["reference.md", "表结构文档"]],
        "outputs": [["reference.md", "表与指标口径说明"]],
        "failures": ["—"],
        "manual": ["—"],
        "related": ["被所有分析 skill 引用"],
    }


def main() -> int:
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("-o", "--output", default=str(DEFAULT_OUT), help="输出 docx 路径")
    args = parser.parse_args()
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build_catalog()
    doc.save(str(out))
    print(f"OK: {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
