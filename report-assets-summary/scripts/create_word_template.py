#!/usr/bin/env python3
"""Create weekly_inspection_report_template.docx (section skeleton for export_report_docx.py)."""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(1)

OUT = Path(__file__).resolve().parents[1] / "templates" / "weekly_inspection_report_template.docx"

PLACEHOLDER = (
    "【自动填充】本节内容由 export_report_docx.py 根据 charts/、tables/、summary.csv 写入。"
    "导出后删除本段说明文字。"
)


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # 封面占位
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Olist 经营巡检周报")
    r.bold = True
    r.font.size = Pt(22)
    doc.add_paragraph()

    for title in [
        "一、执行摘要",
        "二、图表解读",
        "三、数据明细摘要",
        "四、行动建议（待人工确认）",
    ]:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(PLACEHOLDER)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.size = Pt(9)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("附录：结构化总结（summary.csv）", level=1)
    p = doc.add_paragraph(PLACEHOLDER)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.size = Pt(9)

    doc.save(str(OUT))
    print(f"OK: {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
