#!/usr/bin/env python3
"""Export weekly report folder (charts + tables + summary.csv) to Word (.docx)."""
from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("ERROR: 请先安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

TEMPLATE_NAME = "weekly_inspection_report_template.docx"
CHART_ORDER = [
    "delay_trend",
    "delay_top_sellers",
    "bad_review_by_delay_bins",
    "seller_risk_scatter",
    "category_health_matrix",
]
TABLE_ORDER = [
    "weekly_delay_trend",
    "seller_risk_list",
    "problem_categories",
    "delay_by_seller",
    "delay_by_category",
    "delay_by_region",
    "review_by_delay_bins",
]
TABLE_TITLES = {
    "weekly_delay_trend": "周度延迟与体验趋势",
    "seller_risk_list": "卖家风险清单",
    "problem_categories": "问题品类",
    "delay_by_seller": "卖家延迟排名",
    "delay_by_category": "品类延迟排名",
    "delay_by_region": "区域延迟排名",
    "review_by_delay_bins": "延迟分箱与差评率",
}
CHART_TITLES = {
    "delay_trend": "近8周延迟率趋势",
    "delay_top_sellers": "延迟率 TOP10 卖家",
    "bad_review_by_delay_bins": "延迟分箱 vs 差评率",
    "seller_risk_scatter": "卖家风险散点图",
    "category_health_matrix": "品类健康矩阵",
}


def _skill_root() -> Path:
    return Path(__file__).resolve().parents[1]


def get_template_path() -> Path:
    return _skill_root() / "templates" / TEMPLATE_NAME


def _set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def _add_bullet(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        _set_run_font(r0, bold=True)
    r1 = p.add_run(text)
    _set_run_font(r1)
    return p


def _df_to_docx_table(doc: Document, df: pd.DataFrame, max_rows: int = 12):
    if df.empty:
        doc.add_paragraph("（无数据）")
        return
    view = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(view.columns):
        hdr[i].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.4g}" if abs(val) < 1 else f"{val:.2f}"
            else:
                cells[i].text = str(val) if pd.notna(val) else ""
    doc.add_paragraph("")


def _load_summary_csv(report_dir: Path) -> pd.DataFrame | None:
    p = report_dir / "summary.csv"
    if not p.exists():
        return None
    return pd.read_csv(p, encoding="utf-8-sig")


def _infer_period(report_dir: Path, summary: pd.DataFrame | None) -> str:
    if (report_dir / "report.md").exists():
        first = (report_dir / "report.md").read_text(encoding="utf-8").splitlines()
        for line in first[:3]:
            if "|" in line and "周报" in line:
                return line.strip("# ").strip()
    name = report_dir.name.replace("weekly_", "")
    if len(name) == 8 and name.isdigit():
        return f"分析周期截止 {name[:4]}-{name[4:6]}-{name[6:8]}"
    return report_dir.name


def _clear_document_body(doc: Document) -> None:
    """Remove all body content but keep section properties (styles/page setup)."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("sectPr"):
            body.remove(child)


def build_document(report_dir: Path, use_template: bool = True) -> Document:
    report_dir = report_dir.resolve()
    tpl = get_template_path()
    if use_template and tpl.exists():
        doc = Document(str(tpl))
        _clear_document_body(doc)
    else:
        doc = Document()

    summary = _load_summary_csv(report_dir)
    period = _infer_period(report_dir, summary)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Olist 经营巡检周报")
    _set_run_font(r, size=22, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(sub.add_run(period), size=12)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"),
        size=10,
        color=(0x66, 0x66, 0x66),
    )
    doc.add_paragraph("")

    # 一、执行摘要
    _add_heading(doc, "一、执行摘要", 1)
    if summary is not None:
        exec_df = summary[summary["类别"] == "执行摘要"]
        if len(exec_df):
            for _, row in exec_df.iterrows():
                line = f"{row['子类']} · {row['指标名称']}：{row['指标值']}"
                if pd.notna(row.get("说明")) and str(row["说明"]).strip():
                    line += f"（{row['说明']}）"
                _add_bullet(doc, line)
        else:
            doc.add_paragraph("（summary.csv 中无执行摘要行）")
    else:
        doc.add_paragraph("（未找到 summary.csv，请先运行 summarize_report_assets.py）")
    doc.add_paragraph("")

    # 二、图表解读
    _add_heading(doc, "二、图表解读", 1)
    charts_dir = report_dir / "charts"
    if charts_dir.is_dir():
        pngs = {p.stem: p for p in charts_dir.glob("*.png")}
        ordered = [pngs[k] for k in CHART_ORDER if k in pngs]
        ordered += [p for p in sorted(pngs.values()) if p not in ordered]
        for png in ordered:
            _add_heading(doc, CHART_TITLES.get(png.stem, png.stem), 2)
            if summary is not None:
                chart_rows = summary[
                    (summary["类别"] == "图表解读") & (summary["对象ID"].astype(str) == png.stem)
                ]
                if len(chart_rows) and pd.notna(chart_rows.iloc[0].get("说明")):
                    doc.add_paragraph(str(chart_rows.iloc[0]["说明"]))
            try:
                doc.add_picture(str(png), width=Inches(5.8))
            except Exception as e:
                doc.add_paragraph(f"（图片加载失败：{png.name}，{e}）")
            doc.add_paragraph("")
    else:
        doc.add_paragraph("（无 charts 目录）")

    # 三、数据明细
    _add_heading(doc, "三、数据明细摘要", 1)
    tables_dir = report_dir / "tables"
    if tables_dir.is_dir():
        for stem in TABLE_ORDER:
            csv_path = tables_dir / f"{stem}.csv"
            if not csv_path.exists():
                continue
            df = pd.read_csv(csv_path, encoding="utf-8-sig")
            _add_heading(doc, TABLE_TITLES.get(stem, stem), 2)
            if summary is not None:
                detail = summary[
                    (summary["类别"] == "数据明细") & (summary["数据来源"].astype(str).str.contains(stem))
                ]
                for _, row in detail.head(5).iterrows():
                    note = f"{row['指标名称']} {row['指标值']}"
                    if pd.notna(row.get("对象ID")) and str(row["对象ID"]).strip():
                        note += f" · {row['对象ID']}"
                    if pd.notna(row.get("说明")) and str(row["说明"]).strip():
                        note += f"（{row['说明']}）"
                    _add_bullet(doc, note)
            _df_to_docx_table(doc, df)
    else:
        doc.add_paragraph("（无 tables 目录）")

    # 四、行动建议
    _add_heading(doc, "四、行动建议（待人工确认）", 1)
    if summary is not None:
        actions = summary[summary["类别"] == "行动建议"]
        for i, row in actions.iterrows():
            text = str(row.get("指标值", row.get("指标名称", "")))
            manual = row.get("需人工确认", "是")
            suffix = " 【需人工确认】" if str(manual) == "是" else ""
            _add_bullet(doc, f"{text}{suffix}")
    else:
        doc.add_paragraph("1. 对高延迟卖家开展履约复盘。【需人工确认】")
        doc.add_paragraph("2. 对问题品类优化预计送达与商品描述。【需人工确认】")
        doc.add_paragraph("3. 处罚/下架/赔偿须人工审批。【需人工确认】")

    # 附录：完整 summary 表
    if summary is not None:
        doc.add_page_break()
        _add_heading(doc, "附录：结构化总结（summary.csv）", 1)
        _df_to_docx_table(doc, summary, max_rows=50)

    return doc


def export_report_docx(
    report_dir: Path,
    output: Path | None = None,
    use_template: bool = True,
) -> Path:
    report_dir = report_dir.resolve()
    if output is None:
        output = report_dir / f"{report_dir.name}_report.docx"
    doc = build_document(report_dir, use_template=use_template)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description="将报告目录导出为 Word")
    parser.add_argument("report_dir", help="含 charts/ tables/ summary.csv 的目录")
    parser.add_argument("-o", "--output", help="输出 docx 路径，默认 <report_dir>/<name>_report.docx")
    parser.add_argument("--no-template", action="store_true", help="不加载 Word 模板，从零构建")
    args = parser.parse_args()
    try:
        out = Path(args.output) if args.output else None
        path = export_report_docx(
            Path(args.report_dir),
            output=out,
            use_template=not args.no_template,
        )
        print(f"OK: {path}")
        return 0
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
